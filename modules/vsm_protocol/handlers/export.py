import io
import re
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

from modules.vsm_protocol.handlers.human_readable import get_human_message_templates
from utils.helpers import clean_text, format_car_number, format_datetime, safe_str


def get_protocol_message_text(row):
    code = safe_str(row.get("messagecode", ""))
    event_type = safe_str(row.get("event_type", ""))
    message_text = safe_str(row.get("message_text", ""))

    templates = get_human_message_templates(code)

    if event_type == "activation":
        text = templates.get("kurztext_3") or message_text
    elif event_type == "deactivation":
        text = templates.get("kurztext_4") or message_text
    else:
        text = message_text or templates.get("kurztext_2", "")

    text = clean_text(text)

    text = re.sub(
        rf"\bДС\s+{re.escape(code)}\b",
        f"ДС [{code}]",
        text
    )

    if f"[{code}]" in text:
        return text

    return f"[{code}] {text}"


def build_human_readable_entry(row):
    train_id = safe_str(row.get("train_id", ""))
    carnumber = format_car_number(row.get("carnumber", ""))
    timestamp = row.get("timestamp", None)

    timestamp_str = format_datetime(timestamp)
    message_text = get_protocol_message_text(row)

    lines = []

    header_parts = []

    if train_id:
        header_parts.append(f"поезд {train_id}")

    if carnumber:
        header_parts.append(f"вагон {carnumber}")

    if header_parts:
        lines.append(", ".join(header_parts) + ".")

    if timestamp_str and message_text:
        lines.append(f"{timestamp_str} {message_text}.")
    elif timestamp_str:
        lines.append(f"{timestamp_str} Зафиксировано диагностическое сообщение.")

    return "\n".join(lines)


def build_human_readable_protocol_text(timeline_df, train_human_name, dt_from, dt_to):
    lines = [
        "Эксплуатационный протокол",
        "",
        f"Поезд: {train_human_name}",
        f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}",
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
    ]

    if timeline_df is None or timeline_df.empty:
        lines.append("За указанный период диагностические события не обнаружены.")
        return "\n".join(lines)

    for _, row in timeline_df.iterrows():
        lines.append(build_human_readable_entry(row))
        lines.append("")

    return "\n".join(lines).strip()


def export_text_to_docx(protocol_text, file_title="Эксплуатационный протокол"):
    try:
        doc = Document()
        lines = str(protocol_text).splitlines()

        if lines:
            first_line = lines[0].strip()
            if first_line:
                title = doc.add_heading(first_line, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lines = lines[1:]

        for line in lines:
            doc.add_paragraph(clean_text(line)) if line.strip() else doc.add_paragraph("")

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        print(f"EDITED DOCX export error: {e}")
        return None


def get_column_names_map():
    return {
        "train_id": "Номер поезда",
        "carnumber": "Вагон",
        "messagecode": "Код ДС",
        "event_type": "Тип события",
        "timestamp": "Время события",
        "message_text": "Описание",
        "duration_str": "Продолжительность",
        "parsingtime": "Время парсинга",
    }


def prepare_row_for_export(row, col):
    value = row.get(col, "")

    if col == "event_type":
        if value == "activation":
            return "Активация"
        if value == "deactivation":
            return "Деактивация"
        if value == "still_active_marker":
            return "Активно до сих пор"
        return "—"

    if col == "timestamp":
        return format_datetime(value)

    return clean_text(value)


def export_to_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        doc = Document()

        title = doc.add_heading("Эксплуатационный протокол", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Поезд: {train_human_name}")
        doc.add_paragraph(f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}")
        doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph("")

        if not selected_columns:
            selected_columns = [
                "train_id",
                "carnumber",
                "messagecode",
                "event_type",
                "timestamp",
                "message_text",
            ]

        column_names = get_column_names_map()
        selected_cols = [col for col in selected_columns if col in timeline_df.columns]
        headers = [column_names.get(col, col) for col in selected_cols]

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for _, row in timeline_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(selected_cols):
                cells[i].text = str(prepare_row_for_export(row, col))

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    except Exception as e:
        print(f"DOCX export error: {e}")
        return None


def export_human_readable_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        protocol_text = build_human_readable_protocol_text(
            timeline_df,
            train_human_name,
            dt_from,
            dt_to
        )
        return export_text_to_docx(protocol_text)

    except Exception as e:
        print(f"HUMAN DOCX export error: {e}")
        return None


def export_to_xlsx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Эксплуатационный протокол"

        ws["A1"] = f"Поезд: {train_human_name}"
        ws["A2"] = f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}"
        ws["A3"] = f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"

        if not selected_columns:
            selected_columns = [
                "train_id",
                "carnumber",
                "messagecode",
                "event_type",
                "timestamp",
                "message_text",
            ]

        column_names = get_column_names_map()
        selected_cols = [col for col in selected_columns if col in timeline_df.columns]
        headers = [column_names.get(col, col) for col in selected_cols]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        for row_idx, (_, row) in enumerate(timeline_df.iterrows(), 6):
            for col_idx, col in enumerate(selected_cols, 1):
                ws.cell(row=row_idx, column=col_idx, value=prepare_row_for_export(row, col))

        for col in range(1, len(headers) + 1):
            column_letter = ws.cell(row=5, column=col).column_letter
            ws.column_dimensions[column_letter].width = 25

        xlsx_bytes = io.BytesIO()
        wb.save(xlsx_bytes)
        xlsx_bytes.seek(0)
        return xlsx_bytes

    except Exception as e:
        print(f"XLSX export error: {e}")
        return None


def export_to_csv(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    try:
        if not selected_columns:
            selected_columns = [
                "train_id",
                "carnumber",
                "messagecode",
                "event_type",
                "timestamp",
                "message_text",
            ]

        available_columns = [col for col in selected_columns if col in timeline_df.columns]
        df_clean = timeline_df[available_columns].copy()

        if "event_type" in df_clean.columns:
            event_type_map = {
                "activation": "Активация",
                "deactivation": "Деактивация",
                "still_active_marker": "Активно до сих пор",
            }
            df_clean["event_type"] = df_clean["event_type"].map(event_type_map).fillna(df_clean["event_type"])

        if "timestamp" in df_clean.columns:
            df_clean["timestamp"] = df_clean["timestamp"].apply(format_datetime)

        for col in df_clean.columns:
            if df_clean[col].dtype == "object":
                df_clean[col] = df_clean[col].apply(clean_text)

        csv_data = io.StringIO()
        df_clean.to_csv(csv_data, index=False, encoding="utf-8-sig")
        return io.BytesIO(csv_data.getvalue().encode("utf-8-sig"))

    except Exception as e:
        print(f"CSV export error: {e}")
        return None